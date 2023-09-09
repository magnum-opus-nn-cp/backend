from random import randint

from django.conf import settings
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt
from matplotlib.colors import LinearSegmentedColormap

from press_release_nl.processor.models import Text

cmap = LinearSegmentedColormap.from_list("rg", ["w", "y"], N=512)


def to_rgb(vals):
    return f"{hex(int(vals[0]*255))[2:].upper()}{hex(int(vals[1]*255))[2:].upper()}{hex(int(vals[2]*255))[2:].upper()}"


def create_highlighted_document(pk: int, var: str) -> str:
    text = Text.objects.get(pk=pk).description[var]["text"]
    document = Document()
    p = document.add_paragraph()
    for e in [[y.split(">") for y in x.split("<")] for x in text.split("<span ")]:
        if len(e) == 1:
            run = p.add_run()
            run.add_text(e[0][0])
        else:
            for k, v in e:
                if "data-value" in k:
                    if v:
                        run = p.add_run()
                        run.add_text(v)
                        w = float(k.split("=")[-1].replace('"', ""))
                        tag = run._r
                        run.font.size = Pt(11)
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:fill"), to_rgb(cmap(w)))
                        tag.rPr.append(shd)
                else:
                    if v:
                        run = p.add_run()
                        run.add_text(v)
    f = settings.MEDIA_ROOT + f"/docx/{pk}_{randint(1, 1000)}.docx"
    document.save(f)
    return f.replace(settings.MEDIA_ROOT, "/media")
